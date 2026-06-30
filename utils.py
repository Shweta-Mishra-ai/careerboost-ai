"""
utils.py — CareerBoost AI (TOP 1% VERSION)
Fixes:
  1. CV PDF — proper ATS format, no layout bugs, Table for alignment
  2. LinkedIn public data fetch
  3. GitHub repos + bio + languages fetch
  4. Portfolio — top 1% professional with categorized skills
  5. HR Finder + Email templates
"""

import fitz
import docx
import requests
from bs4 import BeautifulSoup
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import (
    SimpleDocTemplate, Paragraph, Spacer, HRFlowable, Table, TableStyle
)
from reportlab.lib.enums import TA_LEFT, TA_CENTER
from reportlab.lib.colors import HexColor
import re
import io
import zipfile
import datetime
from typing import Dict, List, Optional


# ─────────────────────────────────────────────
# FILE PARSING
# ─────────────────────────────────────────────

def parse_pdf(file) -> str:
    try:
        pdf_bytes = file.read() if hasattr(file, 'read') else file
        doc = fitz.open(stream=pdf_bytes, filetype="pdf")
        text = "\n".join(page.get_text() for page in doc)
        doc.close()
        return text.strip()
    except Exception as e:
        raise Exception(f"PDF parse error: {e}")


def parse_docx(file) -> str:
    try:
        # HIDDEN BUG FIX: Streamlit UploadedFile needs BytesIO
        # docx.Document() fails on Streamlit file objects in some versions
        if hasattr(file, 'read'):
            raw = file.read()
            if isinstance(raw, str):
                raw = raw.encode('utf-8')
            d = docx.Document(io.BytesIO(raw))
        else:
            d = docx.Document(file)
        return "\n".join(p.text for p in d.paragraphs).strip()
    except Exception as e:
        raise Exception(f"DOCX parse error: {e}")


def parse_txt(file) -> str:
    try:
        raw = file.read()
        return (raw.decode('utf-8') if isinstance(raw, bytes) else raw).strip()
    except Exception as e:
        raise Exception(f"TXT parse error: {e}")


def parse_cv(file) -> Dict:
    fname = file.name.lower()
    if fname.endswith('.pdf'):
        text = parse_pdf(file)
    elif fname.endswith(('.docx', '.doc')):
        text = parse_docx(file)
    elif fname.endswith('.txt'):
        text = parse_txt(file)
    else:
        raise Exception("Unsupported format. Use PDF, DOCX or TXT.")

    if not text.strip():
        raise Exception("File appears empty or unreadable.")

    try:
        from llm_utils import extract_cv_data_llm
        data = extract_cv_data_llm(text)
        if data and data.get("name"):
            data["raw_text"] = text
            data.setdefault("skills", [])
            edu = data.get("education", [])
            if edu and isinstance(edu[0], dict):
                data["education_structured"] = edu
                data["education"] = [
                    f"{e.get('degree','')} — {e.get('institution','')} ({e.get('year','')})".strip(" —()")
                    for e in edu
                ]
            return data
    except Exception:
        pass

    return {
        'raw_text': text,
        'name': _extract_name(text),
        'email': _extract_email(text),
        'phone': _extract_phone(text),
        'linkedin': _extract_linkedin(text),
        'github': _extract_github(text),
        'skills': _extract_skills(text),
        'experience': _extract_experience(text),
        'education': _extract_education(text),
        'certifications': [],
        'projects': [],
        'current_title': _extract_title(text),
        'years_experience': 0,
        'location': '',
        'summary': '',
    }


# ─────────────────────────────────────────────
# REGEX HELPERS
# ─────────────────────────────────────────────

def _extract_name(text):
    for line in (l.strip() for l in text.split('\n') if l.strip()):
        if 2 <= len(line.split()) <= 4 and '@' not in line and len(line) > 3:
            if not re.match(r'(?i)(experience|education|skills|summary|objective|contact|profile|about|resume|cv)', line):
                return line
    return "Professional"


def _extract_email(text):
    m = re.search(r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,}\b', text)
    return m.group(0) if m else ""


def _extract_phone(text):
    m = re.search(r'(\+?\d{1,3}[\-.\s]?)?(\(?\d{2,4}\)?[\-.\s]?)?\d{3,4}[\-.\s]?\d{4}', text)
    return m.group(0).strip() if m else ""


def _extract_linkedin(text):
    m = re.search(r'linkedin\.com/in/[\w\-]+', text, re.IGNORECASE)
    return f"https://{m.group(0)}" if m else ""


def _extract_github(text):
    m = re.search(r'github\.com/[\w\-]+', text, re.IGNORECASE)
    return f"https://{m.group(0)}" if m else ""


def _extract_title(text):
    titles = ['software engineer','software developer','frontend developer',
              'backend developer','full stack developer','data scientist',
              'data analyst','ml engineer','devops engineer','product manager',
              'web developer','mobile developer']
    lower = text.lower()
    for t in titles:
        if t in lower:
            return t.title()
    return ""


SKILL_KEYWORDS = [
    'python','java','javascript','typescript','go','golang','rust','swift','kotlin',
    'c++','c#','.net','ruby','php','scala','r','matlab',
    'react','angular','vue.js','next.js','nuxt','svelte','html','css','sass',
    'bootstrap','tailwind css','material ui','figma','adobe xd','ui/ux',
    'node.js','express','django','flask','fastapi','spring boot','laravel',
    'graphql','rest api','websocket','grpc','microservices',
    'sql','postgresql','mysql','sqlite','mongodb','redis','elasticsearch',
    'cassandra','dynamodb','snowflake','bigquery',
    'pandas','numpy','matplotlib','seaborn','plotly',
    'machine learning','deep learning','nlp','computer vision','llm',
    'tensorflow','pytorch','scikit-learn','keras','hugging face',
    'langchain','openai','generative ai','rag',
    'aws','azure','gcp','google cloud','ec2','lambda','s3',
    'docker','kubernetes','terraform','ansible','helm',
    'ci/cd','jenkins','github actions','gitlab ci',
    'linux','bash','shell scripting',
    'git','github','gitlab','jira','confluence',
    'power bi','tableau','excel',
    'agile','scrum','kanban',
    'react native','flutter','ios','android',
    'kafka','rabbitmq','celery',
]


def _extract_skills(text):
    lower = text.lower()
    # HIDDEN BUG FIX: kw.title() mangles tech names:
    # "vue.js" -> "Vue.Js", "node.js" -> "Node.Js", "c++" -> "C++", ".net" -> ".Net"
    # Use a proper display map instead
    DISPLAY = {
        # JS/Frontend
        'vue.js':'Vue.js','node.js':'Node.js','next.js':'Next.js',
        'nuxt':'Nuxt.js','express':'Express.js','tailwind css':'Tailwind CSS',
        'react native':'React Native','material ui':'Material UI',
        'adobe xd':'Adobe XD','three.js':'Three.js','webgl':'WebGL',
        # Languages
        '.net':'.NET','c++':'C++','c#':'C#',
        # APIs/Protocols
        'graphql':'GraphQL','grpc':'gRPC','websocket':'WebSocket',
        'rest api':'REST API',
        # DevOps/Cloud
        'ci/cd':'CI/CD','aws':'AWS','gcp':'GCP','linux':'Linux','bash':'Bash',
        'github actions':'GitHub Actions','gitlab ci':'GitLab CI',
        # Data/AI — critical: these are all wrongly cased by .title()
        'tensorflow':'TensorFlow','pytorch':'PyTorch','numpy':'NumPy',
        'pandas':'Pandas','scikit-learn':'Scikit-learn','opencv':'OpenCV',
        'postgresql':'PostgreSQL','mongodb':'MongoDB','mysql':'MySQL',
        'sqlite':'SQLite','redis':'Redis','elasticsearch':'Elasticsearch',
        'dynamodb':'DynamoDB','firebase':'Firebase','supabase':'Supabase',
        # Concepts
        'nlp':'NLP','rag':'RAG','llm':'LLM','ui/ux':'UI/UX',
        'power bi':'Power BI','sql':'SQL',
        # Misc
        'fastapi':'FastAPI','django':'Django','flask':'Flask',
        'kubernetes':'Kubernetes','terraform':'Terraform','ansible':'Ansible',
    }
    seen, result = set(), []
    for kw in SKILL_KEYWORDS:
        if kw in lower and kw not in seen:
            seen.add(kw)
            display = DISPLAY.get(kw, kw.title())
            result.append(display)
    return sorted(result)


def _extract_experience(text):
    """
    BUG 8 FIX: Improved regex experience parser.
    Extracts title, company, duration, and description from common CV formats.
    Used as fallback when no LLM key is available.
    """
    header = re.search(
        r'(?i)(work\s*experience|professional\s*experience|employment\s*history|experience)',
        text
    )
    if not header:
        return []

    start = header.end()
    next_sec = re.search(
        r'(?i)\n(education|skills|certifications|projects|awards|achievements|languages)',
        text[start:]
    )
    chunk = text[start: start + next_sec.start() if next_sec else len(text)]
    lines = [l.strip() for l in chunk.split('\n') if l.strip()]

    # Date pattern: catches "2020 - 2023", "Jan 2020 – Present", "2020–2022" etc
    DATE_PAT = re.compile(
        r'(jan|feb|mar|apr|may|jun|jul|aug|sep|oct|nov|dec|\d{4})[\w\s]*'
        r'[-–—to]+\s*(jan|feb|mar|apr|may|jun|jul|aug|sep|oct|nov|dec|\d{4}|present|current)',
        re.IGNORECASE
    )
    # Bullet/description starters
    BULLET_PAT = re.compile(r'^[•·\-\*>]\s+')

    entries = []
    current = {}

    for line in lines:
        date_match = DATE_PAT.search(line)
        is_bullet = bool(BULLET_PAT.match(line))

        if date_match and not is_bullet:
            # This line likely contains "Title at Company | Date"
            if current:
                entries.append(current)
            duration = date_match.group(0).strip()
            rest = line[:date_match.start()].strip().rstrip('|·-').strip()
            # Split "Software Engineer at Google" or "Software Engineer | Google"
            parts = re.split(r'\s+(?:at|@|\||·|--)\s+', rest, maxsplit=1, flags=re.IGNORECASE)
            current = {
                'title':       parts[0].strip() if parts else rest,
                'company':     parts[1].strip() if len(parts) > 1 else '',
                'duration':    duration,
                'description': '',
            }
        elif is_bullet and current:
            # Bullet point → append to description
            desc_line = BULLET_PAT.sub('', line).strip()
            if current['description']:
                current['description'] += ' ' + desc_line
            else:
                current['description'] = desc_line
        elif not current and len(line) > 10 and not is_bullet:
            # First non-bullet line after section header — likely a job title
            current = {'title': line[:100], 'company': '', 'duration': '', 'description': ''}

        if len(entries) >= 6:
            break

    if current and current not in entries:
        entries.append(current)

    # Trim descriptions
    for e in entries:
        e['description'] = e['description'][:300]

    return entries


def _extract_education(text):
    edu = []
    for kw in [r"bachelor", r"master", r"phd", r"b\.?tech", r"m\.?tech", r"mba", r"b\.?s", r"m\.?s"]:
        for m in re.finditer(kw, text, re.IGNORECASE):
            ctx = text[max(0, m.start()-20): min(len(text), m.end()+100)].strip()
            if ctx not in edu:
                edu.append(ctx)
    return edu or []


# ─────────────────────────────────────────────
# GITHUB DATA FETCH
# ─────────────────────────────────────────────

def get_github_data(github_url: str) -> Dict:
    """Fetch GitHub profile: bio, repos, languages, stars."""
    if not github_url:
        return {}
    match = re.search(r'github\.com/([^/\s]+)', github_url)
    if not match:
        return {}
    username = match.group(1).strip()
    headers = {
        'Accept': 'application/vnd.github.v3+json',
        # HIDDEN BUG FIX: GitHub returns 403 without User-Agent on some endpoints
        'User-Agent': 'CareerBoost-AI/1.0 (github.com/Shweta-Mishra-ai/careerboost-ai)',
    }
    data = {'username': username, 'projects': [], 'bio': '', 'languages': [], 'total_stars': 0}

    try:
        r = requests.get(f"https://api.github.com/users/{username}", headers=headers, timeout=8)
        if r.status_code == 200:
            u = r.json()
            data['bio']          = u.get('bio') or ''
            data['location']     = u.get('location') or ''
            data['blog']         = u.get('blog') or ''
            data['followers']    = u.get('followers', 0)
            data['public_repos'] = u.get('public_repos', 0)
            data['company']      = u.get('company') or ''

        r2 = requests.get(
            f"https://api.github.com/users/{username}/repos?sort=stars&per_page=20",
            headers=headers, timeout=8
        )
        if r2.status_code == 200:
            repos = r2.json()
            lang_count = {}
            total_stars = 0
            for repo in repos:
                if repo.get('fork'):
                    continue
                stars = repo.get('stargazers_count', 0)
                lang  = repo.get('language') or ''
                total_stars += stars
                if lang:
                    lang_count[lang] = lang_count.get(lang, 0) + 1
                desc = repo.get('description') or (f"A {lang} project" if lang else "GitHub repository")
                data['projects'].append({
                    'name': repo.get('name', ''),
                    'description': desc,
                    'language': lang,
                    'stars': stars,
                    'url': repo.get('html_url', ''),
                    'topics': repo.get('topics', []),
                    'updated': repo.get('updated_at', '')[:10],
                })
            data['projects']    = sorted(data['projects'], key=lambda x: x['stars'], reverse=True)[:6]
            data['total_stars'] = total_stars
            data['languages']   = sorted(lang_count, key=lang_count.get, reverse=True)[:6]
    except Exception:
        pass
    return data


def get_github_projects(github_url: str) -> List[Dict]:
    return get_github_data(github_url).get('projects', [])


# ─────────────────────────────────────────────
# LINKEDIN PUBLIC DATA FETCH
# ─────────────────────────────────────────────

def get_linkedin_data(linkedin_url: str) -> Dict:
    """
    BUG 4 FIX: LinkedIn has blocked all public scraping since 2022.
    Requests to linkedin.com return HTTP 999, 403, or a login redirect page
    with no og: meta tags. This function now returns an empty dict immediately
    and lets the caller handle it gracefully rather than silently returning
    empty data after a 10-second timeout.

    To properly integrate LinkedIn data, use one of:
    - LinkedIn API with OAuth (requires company verification)
    - Proxycurl API (~$0.01/call, https://nubela.co/proxycurl)
    - Manual paste: ask user to paste their LinkedIn About section
    """
    if not linkedin_url or 'linkedin.com' not in linkedin_url:
        return {}
    # LinkedIn blocks scraping — return empty with a flag so UI can inform user
    return {'_blocked': True, '_message': 'LinkedIn requires OAuth API access for profile data. Please paste your About section manually.'}


# ─────────────────────────────────────────────
# ENRICH CV WITH GITHUB + LINKEDIN
# ─────────────────────────────────────────────

def enrich_cv_with_external_data(cv_data: Dict, github_url: str = '', linkedin_url: str = '') -> Dict:
    """Merge GitHub + LinkedIn data into cv_data for top 1% output."""
    enriched = cv_data.copy()

    if github_url:
        gh = get_github_data(github_url)
        enriched['github'] = github_url
        if gh.get('bio') and not enriched.get('summary'):
            enriched['github_bio'] = gh['bio']
        if gh.get('location') and not enriched.get('location'):
            enriched['location'] = gh['location']
        existing_skills = set(s.lower() for s in enriched.get('skills', []))
        for lang in gh.get('languages', []):
            if lang.lower() not in existing_skills:
                enriched.setdefault('skills', []).append(lang)
                existing_skills.add(lang.lower())
        existing_proj = {p.get('name','').lower() for p in enriched.get('projects', [])}
        for gp in gh.get('projects', []):
            if gp['name'].lower() not in existing_proj:
                enriched.setdefault('projects', []).append({
                    'name': gp['name'],
                    # HIDDEN BUG FIX: use 'description' not empty string when desc is None
                    'description': gp.get('description') or f"A {gp.get('language','') or 'GitHub'} project.",
                    'url': gp.get('url',''),
                    'language': gp.get('language',''),
                    'stars': gp.get('stars', 0),
                    'topics': gp.get('topics', []),
                })
                existing_proj.add(gp['name'].lower())
        enriched['github_stats'] = {
            'followers': gh.get('followers',0),
            'repos': gh.get('public_repos',0),
            'stars': gh.get('total_stars',0),
            'languages': gh.get('languages',[]),
        }

    if linkedin_url:
        li = get_linkedin_data(linkedin_url)
        enriched['linkedin'] = linkedin_url
        # BUG 4 FIX: LinkedIn is blocked; _blocked flag means no real data was returned
        if not li.get('_blocked'):
            if li.get('name') and enriched.get('name') in ['Professional','',None]:
                enriched['name'] = li['name']
            if li.get('headline') and not enriched.get('current_title'):
                enriched['current_title'] = li['headline']
            if li.get('location') and not enriched.get('location'):
                enriched['location'] = li['location']
            if li.get('about') and not enriched.get('summary'):
                enriched['linkedin_about'] = li['about']
        else:
            # Store the message so the UI can display it to the user
            enriched['_linkedin_blocked_msg'] = li.get('_message', '')

    return enriched


# ─────────────────────────────────────────────
# BUILD CV FROM SCRATCH — Only GitHub / LinkedIn
# No CV upload needed
# ─────────────────────────────────────────────

def build_cv_from_urls(github_url: str = '', linkedin_url: str = '', extra_info: Dict = None) -> Dict:
    """
    Build a complete cv_data dict purely from GitHub + LinkedIn URLs.
    extra_info = manually entered fields: name, email, phone, title, location, years_exp
    Uses LLM to intelligently fill gaps.
    """
    extra = extra_info or {}

    # Base skeleton
    cv_data = {
        'name': extra.get('name', '').strip() or 'Professional',
        'email': extra.get('email', '').strip(),
        'phone': extra.get('phone', '').strip(),
        'location': extra.get('location', '').strip(),
        'current_title': extra.get('title', '').strip(),
        'years_experience': int(extra.get('years_exp', 0) or 0),
        'linkedin': linkedin_url or '',
        'github': github_url or '',
        'skills': [],
        'experience': [],
        'education': [],
        'projects': [],
        'certifications': [],
        'summary': '',
        'raw_text': '',
        'github_bio': '',
        'linkedin_about': '',
    }

    # ── Pull GitHub data ──
    if github_url:
        gh = get_github_data(github_url)

        # Name from GitHub if not provided
        if cv_data['name'] in ['', 'Professional'] and gh.get('username'):
            cv_data['name'] = gh['username'].replace('-', ' ').replace('_', ' ').title()

        cv_data['github_bio'] = gh.get('bio', '') or ''
        if not cv_data['location'] and gh.get('location'):
            cv_data['location'] = gh['location']

        # Skills from languages
        for lang in gh.get('languages', []):
            if lang not in cv_data['skills']:
                cv_data['skills'].append(lang)

        # Projects from repos
        cv_data['projects'] = gh.get('projects', [])

        # GitHub stats
        cv_data['github_stats'] = {
            'followers': gh.get('followers', 0),
            'repos': gh.get('public_repos', 0),
            'stars': gh.get('total_stars', 0),
            'languages': gh.get('languages', []),
        }

        # Infer more skills from project topics + descriptions
        all_text = ' '.join(
            p.get('description', '') + ' ' + ' '.join(p.get('topics', []))
            for p in cv_data['projects']
        ).lower()
        for kw in SKILL_KEYWORDS:
            if kw in all_text and kw.title() not in cv_data['skills']:
                cv_data['skills'].append(kw.title())

    # ── Pull LinkedIn data ──
    if linkedin_url:
        li = get_linkedin_data(linkedin_url)
        if li.get('name') and cv_data['name'] in ['', 'Professional']:
            cv_data['name'] = li['name']
        if li.get('headline') and not cv_data['current_title']:
            cv_data['current_title'] = li['headline']
        if li.get('location') and not cv_data['location']:
            cv_data['location'] = li['location']
        if li.get('about'):
            cv_data['linkedin_about'] = li['about']
            # Extract skills from about text
            about_lower = li['about'].lower()
            for kw in SKILL_KEYWORDS:
                if kw in about_lower and kw.title() not in cv_data['skills']:
                    cv_data['skills'].append(kw.title())

    # ── Use LLM to generate experience + enrich everything ──
    if cv_data['projects'] or cv_data['github_bio'] or cv_data['linkedin_about']:
        raw_context = f"""
Name: {cv_data['name']}
Title: {cv_data['current_title']}
GitHub Bio: {cv_data.get('github_bio', '')}
LinkedIn About: {cv_data.get('linkedin_about', '')}
GitHub Projects: {', '.join(p['name'] + ' (' + p.get('language','') + ')' for p in cv_data['projects'][:5])}
GitHub Languages: {', '.join(cv_data['github_stats'].get('languages', []) if cv_data.get('github_stats') else [])}
Skills detected: {', '.join(cv_data['skills'][:15])}
Years experience: {cv_data['years_experience']}
"""
        cv_data['raw_text'] = raw_context

        try:
            from llm_utils import enrich_from_github_llm
            enriched = enrich_from_github_llm(cv_data)
            if enriched:
                # Merge LLM enriched data carefully
                if enriched.get('current_title') and not cv_data['current_title']:
                    cv_data['current_title'] = enriched['current_title']
                if enriched.get('experience'):
                    cv_data['experience'] = enriched['experience']
                if enriched.get('skills'):
                    merged = list({s.lower(): s for s in cv_data['skills'] + enriched['skills']}.values())
                    cv_data['skills'] = merged[:30]
                if enriched.get('education') and not cv_data['education']:
                    cv_data['education'] = enriched['education']
                if enriched.get('summary'):
                    cv_data['summary'] = enriched['summary']
        except Exception:
            pass

    return cv_data


# ─────────────────────────────────────────────
# ATS ANALYSIS
# ─────────────────────────────────────────────

def analyze_ats(cv_data: Dict, job_description: str) -> Dict:
    try:
        from llm_utils import analyze_ats_llm
        result = analyze_ats_llm(cv_data.get('raw_text',''), job_description)
        if result and result.get('score') is not None:
            return result
    except Exception:
        pass

    cv_lower  = cv_data.get('raw_text','').lower()
    job_lower = job_description.lower()
    job_kws   = [kw for kw in SKILL_KEYWORDS if kw in job_lower]
    matched   = [kw for kw in job_kws if kw in cv_lower]
    missing   = [kw for kw in job_kws if kw not in cv_lower]

    score = round((len(matched)/len(job_kws))*100) if job_kws else 65
    if not job_kws:
        if 'project' in cv_lower: score += 5
        if '%' in cv_lower: score += 10
        if len(cv_data.get('skills',[])) > 5: score += 10
    score = max(0, min(100, score))

    strengths = []
    if len(matched) > 5: strengths.append("Strong keyword alignment with job requirements.")
    if 'project' in cv_lower: strengths.append("Projects section detected — great credibility signal.")
    if any(v in cv_lower for v in ['%','increased','reduced','improved']):
        strengths.append("Quantified achievements present — excellent for recruiters.")
    if cv_data.get('github'): strengths.append("GitHub linked — shows real technical portfolio.")

    return {
        'score': score,
        'matched_skills': [s.title() for s in matched],
        'missing_skills': [s.title() for s in missing[:10]],
        'semantic_gaps': [],
        'strengths': strengths,
        'tips': _generate_tips(cv_data, missing, cv_lower),
        'experience_match': score,
        'keyword_density': 'high' if score>75 else 'medium' if score>50 else 'low',
    }


def _generate_tips(cv_data, missing, cv_lower):
    resources = {
        'docker': 'Play With Docker (free)', 'aws': 'AWS Free Tier + freeCodeCamp',
        'kubernetes': 'Kubernetes.io tutorial', 'python': 'Python.org + freeCodeCamp',
        'javascript': 'JavaScript.info (free)', 'react': 'react.dev official',
        'sql': 'SQLBolt (free)', 'machine learning': 'Andrew Ng Coursera (audit)',
    }
    tips = []
    for skill in missing[:4]:
        res = resources.get(skill.lower(), 'YouTube + freeCodeCamp')
        tips.append(f"Add <strong>{skill.title()}</strong> — Learn: {res}")
    if 'project' not in cv_lower:
        tips.append("Add a <strong>Projects section</strong> with GitHub links")
    if not any(v in cv_lower for v in ['%','increased','reduced','grew']):
        tips.append("<strong>Quantify achievements</strong> — e.g. 'Reduced load time by 40%'")
    if 'certification' not in cv_lower:
        tips.append("Add <strong>free certifications</strong> — Google, AWS, Meta on Coursera")
    return tips[:8]


def generate_skills_roadmap(missing_skills, target_role=""):
    try:
        from llm_utils import generate_roadmap_llm
        result = generate_roadmap_llm(missing_skills, target_role)
        if result and len(result) > 50:
            return result
    except Exception:
        pass
    # HIDDEN BUG FIX: was returning None on failure -> blank roadmap tab
    # Now returns a useful fallback markdown
    db = {
        'python': {'weeks':'3-5','r':['Python.org','freeCodeCamp']},
        'react':  {'weeks':'3-4','r':['react.dev','freeCodeCamp React']},
        'docker': {'weeks':'2-3','r':['Docker Docs','Play With Docker']},
        'aws':    {'weeks':'6-8','r':['AWS Free Tier','freeCodeCamp AWS']},
        'sql':    {'weeks':'2-3','r':['SQLBolt','W3Schools SQL']},
    }
    md = f"# Roadmap\n*{datetime.datetime.now().strftime('%B %d, %Y')}*\n\n"
    for i, skill in enumerate(missing_skills[:8], 1):
        info = db.get(skill.lower(), {'weeks':'2-4','r':['YouTube','freeCodeCamp']})
        md += f"## {i}. {skill.title()}\n⏱️ {info['weeks']} weeks\n"
        for r in info['r']: md += f"  - {r}\n"
        md += "\n"
    return md


# ─────────────────────────────────────────────
# CV PDF — TOP 1% ATS FORMAT (FIXED)
# ─────────────────────────────────────────────

def generate_optimized_cv(cv_data: Dict, job_description: str = None, template: str = "Modern") -> bytes:
    buf = io.BytesIO()
    is_harvard = template.lower() == "harvard"

    doc = SimpleDocTemplate(
        buf, pagesize=letter,
        topMargin=0.55*inch, bottomMargin=0.55*inch,
        leftMargin=0.65*inch, rightMargin=0.65*inch,
    )
    styles = getSampleStyleSheet()

    # Colors
    if is_harvard:
        CA = HexColor('#A51C30')  # Harvard Crimson
        CN = HexColor('#000000')
        CB = HexColor('#1a1a1a')
        CM = HexColor('#555555')
        FH = 'Times-Bold'
        FB = 'Times-Roman'
        FI = 'Times-Italic'
        NS = 22; NA = TA_CENTER
    else:
        CA = HexColor('#1a56db')  # Blue
        CN = HexColor('#0d1117')
        CB = HexColor('#24292f')
        CM = HexColor('#57606a')
        FH = 'Helvetica-Bold'
        FB = 'Helvetica'
        FI = 'Helvetica-Oblique'
        NS = 24; NA = TA_LEFT

    def sty(name, **kwargs):
        return ParagraphStyle(name, parent=styles['Normal'], **kwargs)

    s_name    = sty('nm',  fontSize=NS, leading=NS+4, textColor=CN, alignment=NA, fontName=FH, spaceAfter=2)
    s_title   = sty('jt',  fontSize=11, leading=14, textColor=CA, alignment=TA_CENTER, fontName=FH, spaceAfter=2)
    s_contact = sty('ct',  fontSize=8.5, leading=12, textColor=CM, alignment=TA_CENTER, fontName=FB, spaceAfter=3)
    s_section = sty('sec', fontSize=10, leading=13, textColor=CA, alignment=TA_CENTER if is_harvard else TA_LEFT,
                    fontName=FH, spaceBefore=9, spaceAfter=1)
    s_body    = sty('bd',  fontSize=9.5, leading=13.5, textColor=CB, fontName=FB, spaceAfter=2)
    s_role    = sty('rl',  fontSize=10, leading=13, textColor=CB, fontName=FH, spaceAfter=1)
    s_company = sty('co',  fontSize=9, leading=12, textColor=CA, fontName=FB, spaceAfter=1)
    s_date    = sty('dt',  fontSize=8.5, leading=11, textColor=CM, fontName=FI, spaceAfter=2)
    s_bullet  = sty('bu',  fontSize=9, leading=13, textColor=CB, fontName=FB, leftIndent=10, spaceAfter=1.5)
    s_date_r  = sty('dtr', fontSize=8.5, leading=11, textColor=CM, fontName=FI, alignment=2)  # right align

    def rule(thick=0.6):
        return HRFlowable(width='100%', color=CA, thickness=thick, spaceAfter=4, spaceBefore=0)

    def section(title):
        return [Paragraph(title.upper(), s_section), rule()]

    story = []

    # ─ HEADER ─
    name = (cv_data.get('name') or 'Professional').strip()
    jtitle = (cv_data.get('current_title') or '').strip()
    story.append(Paragraph(name.upper() if is_harvard else name, s_name))
    if jtitle:
        story.append(Paragraph(jtitle, s_title))

    # Contact row 1: email | phone | location
    r1 = [x for x in [
        cv_data.get('email'), cv_data.get('phone'), cv_data.get('location')
    ] if x]
    if r1:
        story.append(Paragraph('  |  '.join(r1), s_contact))

    # Contact row 2: linkedin | github
    r2_parts = []
    if cv_data.get('linkedin'):
        r2_parts.append(cv_data['linkedin'].replace('https://','').replace('http://',''))
    if cv_data.get('github'):
        r2_parts.append(cv_data['github'].replace('https://','').replace('http://',''))
    if r2_parts:
        story.append(Paragraph('  |  '.join(r2_parts), s_contact))

    story.append(HRFlowable(width='100%', color=CA, thickness=1.5, spaceAfter=5, spaceBefore=4))

    # ─ SUMMARY ─
    # BUG 5 FIX: Use pre-generated summary from cv_data if available (set by caller),
    # otherwise fall back gracefully — no inline LLM call inside PDF generation.
    story.extend(section("PROFESSIONAL SUMMARY"))
    summary = cv_data.get('_pre_generated_summary', '')
    if not summary:
        # Build a decent fallback without an LLM call
        skills_str = ', '.join(cv_data.get('skills', [])[:5]) or 'various technologies'
        t = cv_data.get('current_title', 'professional')
        fallback_text = cv_data.get('linkedin_about', '') or cv_data.get('github_bio', '') or cv_data.get('summary', '')
        summary = fallback_text[:300] if fallback_text else (
            f"Results-driven {t} with expertise in {skills_str}. "
            "Passionate about building scalable solutions and delivering measurable impact."
        )
    story.append(Paragraph(summary, s_body))

    # ─ SKILLS ─
    skills = cv_data.get('skills', [])
    if skills:
        story.extend(section("TECHNICAL SKILLS"))
        if job_description:
            jd_l = job_description.lower()
            matched = [s for s in skills if s.lower() in jd_l]
            rest    = [s for s in skills if s.lower() not in jd_l]
            ordered = matched + rest
        else:
            ordered = skills
        # Rows of 6
        for i in range(0, min(len(ordered), 24), 6):
            chunk = ordered[i:i+6]
            story.append(Paragraph('  •  '.join(chunk), s_body))

    # ─ EXPERIENCE ─
    experience = cv_data.get('experience', [])
    if experience:
        story.extend(section("PROFESSIONAL EXPERIENCE"))
        for exp in experience[:5]:
            role    = (exp.get('title') or '').strip()
            company = (exp.get('company') or '').strip()
            dur     = (exp.get('duration') or '').strip()
            desc    = (exp.get('description') or '').strip()
            if not role:
                continue

            # Role + date on same line using Table
            if dur:
                tbl = Table(
                    [[Paragraph(f"<b>{role}</b>", s_role), Paragraph(dur, s_date_r)]],
                    colWidths=[4.4*inch, 2.3*inch]
                )
                tbl.setStyle(TableStyle([
                    ('VALIGN',(0,0),(-1,-1),'TOP'),
                    ('LEFTPADDING',(0,0),(-1,-1),0), ('RIGHTPADDING',(0,0),(-1,-1),0),
                    ('TOPPADDING',(0,0),(-1,-1),0),  ('BOTTOMPADDING',(0,0),(-1,-1),1),
                ]))
                story.append(tbl)
            else:
                story.append(Paragraph(f"<b>{role}</b>", s_role))

            if company:
                story.append(Paragraph(company, s_company))

            if desc:
                bullets = [b.strip() for b in re.split(r'[\n;•\-]', desc) if len(b.strip()) > 10]
                if bullets:
                    for b in bullets[:4]:
                        story.append(Paragraph(f"• {b}", s_bullet))
                else:
                    story.append(Paragraph(f"• {desc[:300]}", s_bullet))
            story.append(Spacer(1, 3))

    # ─ PROJECTS ─
    projects = cv_data.get('projects', [])
    if projects:
        story.extend(section("PROJECTS"))
        for proj in projects[:5]:
            pname  = (proj.get('name') or '').strip()
            pdesc  = (proj.get('description') or '').strip()
            plang  = (proj.get('language') or '').strip()
            purl   = (proj.get('url') or '').strip()
            pstars = proj.get('stars', 0) or 0
            if not pname:
                continue
            meta = []
            if plang:  meta.append(plang)
            if pstars: meta.append(f"★ {pstars}")
            name_line = f"<b>{pname}</b>" + (f"  <font size='8' color='#888888'>[{', '.join(meta)}]</font>" if meta else "")
            story.append(Paragraph(name_line, s_role))
            if pdesc:
                story.append(Paragraph(f"• {pdesc[:200]}", s_bullet))
            if purl:
                story.append(Paragraph(f"<font size='8' color='#1a56db'>{purl.replace('https://','')}</font>", s_date))
            story.append(Spacer(1, 3))

    # ─ EDUCATION ─
    education = cv_data.get('education', [])
    if education:
        story.extend(section("EDUCATION"))
        edu_structured = cv_data.get('education_structured', [])
        if edu_structured:
            for e in edu_structured[:3]:
                degree = (e.get('degree') or '').strip()
                inst   = (e.get('institution') or '').strip()
                year   = (e.get('year') or '').strip()
                if degree or inst:
                    if year:
                        tbl = Table(
                            [[Paragraph(f"<b>{degree}</b>", s_role), Paragraph(year, s_date_r)]],
                            colWidths=[4.8*inch, 1.9*inch]
                        )
                        tbl.setStyle(TableStyle([
                            ('VALIGN',(0,0),(-1,-1),'TOP'),
                            ('LEFTPADDING',(0,0),(-1,-1),0),('RIGHTPADDING',(0,0),(-1,-1),0),
                            ('TOPPADDING',(0,0),(-1,-1),0),('BOTTOMPADDING',(0,0),(-1,-1),1),
                        ]))
                        story.append(tbl)
                    else:
                        story.append(Paragraph(f"<b>{degree}</b>", s_role))
                    if inst:
                        story.append(Paragraph(inst, s_company))
                    story.append(Spacer(1, 2))
        else:
            for e in education[:3]:
                story.append(Paragraph(f"• {str(e)[:120]}", s_body))

    # ─ CERTIFICATIONS ─
    certs = cv_data.get('certifications', [])
    if certs:
        story.extend(section("CERTIFICATIONS"))
        story.append(Paragraph('  |  '.join(str(c) for c in certs[:6]), s_body))

    doc.build(story)
    buf.seek(0)
    return buf.getvalue()


# ─────────────────────────────────────────────
# HR FINDER
# ─────────────────────────────────────────────

def find_hr_contacts(company_name: str, role: str = "") -> List[Dict]:
    """Find HR/Recruiter contacts at a company using free public methods."""
    if not company_name:
        return []
    company_slug   = re.sub(r'[^a-z0-9]', '', company_name.lower())
    company_domain = f"{company_slug}.com"
    contacts = []

    # LinkedIn search URLs
    hr_titles = ["Technical Recruiter", "HR Manager", "Talent Acquisition", "People Operations", "Hiring Manager"]
    for title in hr_titles[:4]:
        q = f"{title} {company_name}".replace(' ', '%20')
        contacts.append({
            'type': 'linkedin_search',
            'name': f"{title} at {company_name}",
            'title': title,
            'company': company_name,
            'linkedin_search_url': f"https://www.linkedin.com/search/results/people/?keywords={q}&origin=GLOBAL_SEARCH_HEADER",
            'action': 'Open link → connect + message',
        })

    # Common email patterns
    email_patterns = [
        f"hr@{company_domain}", f"careers@{company_domain}",
        f"talent@{company_domain}", f"recruiting@{company_domain}",
        f"jobs@{company_domain}", f"hiring@{company_domain}",
    ]
    for ep in email_patterns:
        contacts.append({
            'type': 'email_pattern',
            'email': ep,
            'company': company_name,
            'note': 'Common HR email pattern — verify with tools like Hunter.io (free)',
        })

    return contacts


def generate_hr_email_templates(cv_data: Dict, company: str, role: str, hr_name: str = "Hiring Manager") -> Dict:
    """Generate cold outreach, follow-up x2, and thank you email templates."""
    name   = cv_data.get('name', 'Candidate')
    title  = cv_data.get('current_title', 'Professional')
    skills = ', '.join(cv_data.get('skills', [])[:5])
    email  = cv_data.get('email', '')
    today  = datetime.date.today().strftime("%B %d, %Y")

    try:
        from llm_utils import generate_hr_emails_llm
        return generate_hr_emails_llm(cv_data, company, role, hr_name)
    except Exception:
        pass

    cold = f"""Subject: {role} — {name} ({title})

Dear {hr_name},

I hope this finds you well. I'm a {title} with expertise in {skills}, and I'm very interested in the {role} position at {company}.

What draws me specifically to {company} is your reputation for technical excellence and innovation. I believe my background aligns strongly with what your team needs — I've been building [mention a specific relevant achievement] and would love to bring that experience to your team.

My CV is attached for your review. I'd welcome a 15-minute call at your convenience.

Thank you for your time.

Best regards,
{name}
{email} | {today}"""

    followup_1 = f"""Subject: Following Up — {role} Application — {name}

Dear {hr_name},

I'm following up on my application for the {role} position I sent last week.

I remain genuinely excited about the opportunity to join {company}. My experience with {skills} maps closely to what you're building, and I'd love to discuss further.

Happy to provide any additional information you might need.

Thank you,
{name}
{email}"""

    followup_2 = f"""Subject: Quick Follow-Up — {role} — {name}

Hi {hr_name},

I wanted to reach out one more time about the {role} role at {company}.

I'm still very interested and believe I could add immediate value. If this isn't the right timing, I'd also be glad to be considered for future openings.

Thanks for your consideration either way.

{name}
{email}"""

    thank_you = f"""Subject: Thank You — {role} Interview — {name}

Dear {hr_name},

Thank you for taking the time to interview me today for the {role} position.

I really enjoyed our conversation — especially learning about [specific topic from interview]. It reinforced my excitement about joining {company}.

I'm confident my experience in {skills} will help me contribute meaningfully from day one. Please feel free to reach out if you need anything else from my side.

Looking forward to next steps!

Warm regards,
{name}
{email}"""

    # HIDDEN BUG FIX: Return as {subject, body} dicts to match LLM output format
    # Old code returned flat strings which caused KeyError in UI email renderer
    def _split(s):
        lines = s.split("\n", 1)
        subj = lines[0].replace("Subject:","").strip()
        body = lines[1].strip() if len(lines) > 1 else s
        return {"subject": subj, "body": body}

    return {
        "cold_email":  _split(cold),
        "follow_up_1": _split(followup_1),
        "follow_up_2": _split(followup_2),
        "thank_you":   _split(thank_you),
    }


# ─────────────────────────────────────────────
# PORTFOLIO HTML — TOP 1%
# ─────────────────────────────────────────────

def generate_portfolio(cv_data: Dict) -> bytes:
    """
    Generate a stunning, production-quality portfolio HTML.
    Design: Dark glassmorphism, animated gradients, scroll reveals,
    project cards with hover effects, categorized skill grid, timeline experience.
    """
    import io, zipfile, datetime as dt

    name     = (cv_data.get("name") or "Professional").strip()
    email    = (cv_data.get("email") or "").strip()
    phone    = (cv_data.get("phone") or "").strip()
    linkedin = (cv_data.get("linkedin") or "").strip()
    github   = (cv_data.get("github") or "").strip()
    title    = (cv_data.get("current_title") or "Software Developer").strip()
    skills   = cv_data.get("skills", []) or []
    exp      = cv_data.get("experience", []) or []
    edu      = cv_data.get("education", []) or []
    projects = cv_data.get("projects", []) or []
    certs    = cv_data.get("certifications", []) or []
    years    = cv_data.get("years_experience", 0) or 0
    gh_stats = cv_data.get("github_stats", {}) or {}
    summary  = (cv_data.get("summary") or cv_data.get("linkedin_about") or cv_data.get("github_bio") or "").strip()
    about    = summary or (
        f"I'm a {title} passionate about building clean, scalable software. "
        f"I specialize in {chr(44).join(skills[:3]) if skills else 'modern technologies'} "
        "and thrive in collaborative environments."
    )

    initials   = "".join(w[0].upper() for w in name.split()[:2]) if name else "CB"
    stat_num   = f"{years}+" if years else str(len(exp))
    stat_lbl   = "Years Exp" if years else "Roles"
    year_now   = dt.datetime.now().year

    # ── Skill grid by category ──────────────────────────────
    CATS = {
        "Languages":    ["Python","Java","JavaScript","TypeScript","Go","Golang","Rust","Swift","Kotlin","C++","C#","Ruby","PHP","Scala","Dart"],
        "Frontend":     ["React","Angular","Vue.js","Next.js","Svelte","HTML","CSS","Tailwind CSS","Bootstrap","Figma","WebGL","Three.js"],
        "Backend":      ["Node.js","Express","Django","Flask","FastAPI","Spring Boot","GraphQL","REST API","Microservices","gRPC"],
        "Data & AI":    ["Machine Learning","Deep Learning","Pandas","NumPy","TensorFlow","PyTorch","NLP","SQL","MongoDB","PostgreSQL","Redis"],
        "Cloud/DevOps": ["Docker","Kubernetes","AWS","Azure","GCP","Terraform","CI/CD","Jenkins","GitHub Actions","Linux","Nginx"],
        "Tools":        ["Git","GitHub","Jira","Agile","Scrum","Figma","Tableau","Power BI","Postman"],
    }
    user_lower   = {s.lower() for s in skills}
    uncategorized = list(skills)
    skill_grid   = ""
    for cat, cat_skills in CATS.items():
        matched = [s for s in cat_skills if s.lower() in user_lower]
        if matched:
            uncategorized = [s for s in uncategorized if s.lower() not in {m.lower() for m in matched}]
            pills = "".join(f'<span class="sp">{s}</span>' for s in matched)
            skill_grid += f'<div class="sg-row"><div class="sg-cat">{cat}</div><div class="sg-pills">{pills}</div></div>'
    if not skill_grid:
        pills = "".join(f'<span class="sp">{s}</span>' for s in skills[:24])
        skill_grid = f'<div class="sg-row"><div class="sg-cat">Technologies</div><div class="sg-pills">{pills}</div></div>'
    elif uncategorized[:6]:
        pills = "".join(f'<span class="sp">{s}</span>' for s in uncategorized[:6])
        skill_grid += f'<div class="sg-row"><div class="sg-cat">Other</div><div class="sg-pills">{pills}</div></div>'

    # ── Experience timeline ──────────────────────────────────
    exp_html = ""
    for i, e in enumerate(exp[:6]):
        role    = (e.get("title") or "").strip()
        company = (e.get("company") or "").strip()
        dur     = (e.get("duration") or "").strip()
        desc    = (e.get("description") or "").strip()
        if not role:
            continue
        dot_cls = 'tl-dot current' if i == 0 else 'tl-dot'
        exp_html += f"""
        <div class="tl-item rv">
          <div class="tl-lc"><div class="{dot_cls}"></div><div class="tl-line"></div></div>
          <div class="tl-body">
            <div class="tl-hdr">
              <div><div class="tl-role">{role}</div>{"<div class='tl-company'>"+company+"</div>" if company else ""}</div>
              {"<div class='tl-dur'>"+dur+"</div>" if dur else ""}
            </div>
            {"<p class='tl-desc'>"+desc[:260]+"</p>" if desc else ""}
          </div>
        </div>"""

    # ── Project cards ────────────────────────────────────────
    GRADS = [("#5b6af9","#f95b8d"),("#5bf9c0","#5b6af9"),("#f9c85b","#f95b8d"),
             ("#f95b8d","#5bf9c0"),("#a78bfa","#5b6af9"),("#5bf9c0","#f9c85b")]
    proj_html = ""
    for i, p in enumerate(projects[:6]):
        pname   = (p.get("name") or "").strip()
        pdesc   = (p.get("description") or "").strip()
        plang   = (p.get("language") or "").strip()
        purl    = (p.get("url") or "").strip()
        pstars  = p.get("stars", 0) or 0
        ptopics = p.get("topics", []) or p.get("tech", []) or []
        if not pname:
            continue
        g1, g2    = GRADS[i % len(GRADS)]
        topics_h  = "".join(f'<span class="pt">{t}</span>' for t in ptopics[:4])
        proj_html += f"""
        <div class="pc rv">
          <div class="pc-glow" style="background:linear-gradient(135deg,{g1}22,{g2}11)"></div>
          <div class="pc-in">
            <div class="pc-top">
              <div class="pc-av" style="background:linear-gradient(135deg,{g1},{g2})">{pname[0].upper()}</div>
              <div class="pc-bgs">{"<span class='pb-lang'>"+plang+"</span>" if plang else ""}{"<span class='pb-st'>★ "+str(pstars)+"</span>" if pstars else ""}</div>
            </div>
            <div class="pc-name">{pname}</div>
            <p class="pc-desc">{pdesc[:180]}</p>
            {"<div class='pc-topics'>"+topics_h+"</div>" if topics_h else ""}
            <div class="pc-foot">{"<a href='"+purl+"' target='_blank' class='pc-link'>View project <span>→</span></a>" if purl else "<span></span>"}</div>
          </div>
        </div>"""

    # ── Education ────────────────────────────────────────────
    edu_html = ""
    for e in edu[:3]:
        if isinstance(e, dict):
            deg  = (e.get("degree") or "").strip()
            inst = (e.get("institution") or "").strip()
            yr   = (e.get("year") or "").strip()
            if deg or inst:
                edu_html += f'<div class="edu-card rv"><div class="edu-icon">🎓</div><div class="edu-info"><div class="edu-deg">{deg or inst}</div>{"<div class='edu-inst'>"+inst+(" · "+yr if yr else "")+"</div>" if inst and deg else ""}</div></div>'
        else:
            edu_html += f'<div class="edu-card rv"><div class="edu-icon">🎓</div><div class="edu-info"><div class="edu-deg">{str(e)[:80]}</div></div></div>'

    # ── Certs, GitHub stats, Socials ──────────────────────────
    certs_html  = "".join(f'<span class="cert">{c}</span>' for c in certs[:8])
    ghstats_h   = ""
    if gh_stats:
        ghstats_h = (f'<div class="gh-stats">'
            f'<div class="gh-s"><span class="gh-n">{gh_stats.get("repos",0)}</span><span class="gh-l">Repos</span></div>'
            f'<div class="gh-s"><span class="gh-n">{gh_stats.get("stars",0)}</span><span class="gh-l">Stars</span></div>'
            f'<div class="gh-s"><span class="gh-n">{gh_stats.get("followers",0)}</span><span class="gh-l">Followers</span></div>'
            f'</div>')

    soc = ""
    LI_SVG = '<svg width="15" height="15" viewBox="0 0 24 24" fill="currentColor"><path d="M20.447 20.452h-3.554v-5.569c0-1.328-.027-3.037-1.852-3.037-1.853 0-2.136 1.445-2.136 2.939v5.667H9.351V9h3.414v1.561h.046c.477-.9 1.637-1.85 3.37-1.85 3.601 0 4.267 2.37 4.267 5.455v6.286zM5.337 7.433a2.062 2.062 0 01-2.063-2.065 2.064 2.064 0 112.063 2.065zm1.782 13.019H3.555V9h3.564v11.452zM22.225 0H1.771C.792 0 0 .774 0 1.729v20.542C0 23.227.792 24 1.771 24h20.451C23.2 24 24 23.227 24 22.271V1.729C24 .774 23.2 0 22.222 0h.003z"/></svg>'
    GH_SVG = '<svg width="15" height="15" viewBox="0 0 24 24" fill="currentColor"><path d="M12 0C5.374 0 0 5.373 0 12c0 5.302 3.438 9.8 8.207 11.387.599.111.793-.261.793-.577v-2.234c-3.338.726-4.033-1.416-4.033-1.416-.546-1.387-1.333-1.756-1.333-1.756-1.089-.745.083-.729.083-.729 1.205.084 1.839 1.237 1.839 1.237 1.07 1.834 2.807 1.304 3.492.997.107-.775.418-1.305.762-1.604-2.665-.305-5.467-1.334-5.467-5.931 0-1.311.469-2.381 1.236-3.221-.124-.303-.535-1.524.117-3.176 0 0 1.008-.322 3.301 1.23A11.509 11.509 0 0112 5.803c1.02.005 2.047.138 3.006.404 2.291-1.552 3.297-1.23 3.297-1.23.653 1.653.242 2.874.118 3.176.77.84 1.235 1.911 1.235 3.221 0 4.609-2.807 5.624-5.479 5.921.43.372.823 1.102.823 2.222v3.293c0 .319.192.694.801.576C20.566 21.797 24 17.3 24 12c0-6.627-5.373-12-12-12z"/></svg>'
    if linkedin: soc += f'<a href="{linkedin}" target="_blank" class="sb sb-li">{LI_SVG} LinkedIn</a>'
    if github:   soc += f'<a href="{github}" target="_blank" class="sb sb-gh">{GH_SVG} GitHub</a>'
    if email:    soc += f'<a href="mailto:{email}" class="sb sb-em">✉ Email Me</a>'

    nav_projects = "<a href='#projects'>Projects</a>" if projects else ""
    hire_btn     = f'<a href="mailto:{email}" class="nav-cta">Hire Me</a>' if email else ""
    about_contact = ""
    if email or phone:
        about_contact = f'<div><p class="about-cl">Contact</p>{"<p class='about-cv'>✉ "+email+"</p>" if email else ""}{"<p class='about-cv'>📱 "+phone+"</p>" if phone else ""}</div>'

    html = f"""<!DOCTYPE html>
<html lang="en">
<head>
<meta charset="UTF-8"/>
<meta name="viewport" content="width=device-width,initial-scale=1"/>
<meta name="description" content="{name} — {title} · Portfolio"/>
<title>{name} · {title}</title>
<link rel="preconnect" href="https://fonts.googleapis.com"/>
<link href="https://fonts.googleapis.com/css2?family=Plus+Jakarta+Sans:wght@300;400;500;600;700;800&display=swap" rel="stylesheet"/>
<style>
*,*::before,*::after{{box-sizing:border-box;margin:0;padding:0}}
:root{{
  --bg:#07070f;--s1:#0b0b17;--card:#10101e;
  --bd:rgba(255,255,255,.07);--bd2:rgba(255,255,255,.13);
  --a:#6366f1;--a2:#ec4899;--a3:#10b981;--a4:#f59e0b;
  --t:#f0f0ff;--m:#6b6b9a;--ms:#9898c0;
  --font:"Plus Jakarta Sans",system-ui,sans-serif;--r:14px;--r2:20px;
}}
html{{scroll-behavior:smooth}}
body{{font-family:var(--font);background:var(--bg);color:var(--t);line-height:1.65;overflow-x:hidden;-webkit-font-smoothing:antialiased}}
.orb{{position:fixed;border-radius:50%;filter:blur(130px);opacity:.08;pointer-events:none;z-index:0}}
.o1{{width:900px;height:900px;background:radial-gradient(circle,var(--a),transparent);top:-350px;left:-250px;animation:drift 32s infinite alternate ease-in-out}}
.o2{{width:700px;height:700px;background:radial-gradient(circle,var(--a2),transparent);bottom:-250px;right:-200px;animation:drift 26s infinite alternate-reverse ease-in-out}}
.o3{{width:500px;height:500px;background:radial-gradient(circle,var(--a3),transparent);top:40%;left:45%;animation:drift 40s infinite alternate ease-in-out}}
@keyframes drift{{0%{{transform:translate(0,0) scale(1)}}100%{{transform:translate(28px,18px) scale(1.07)}}}}
nav{{position:fixed;top:0;left:0;right:0;z-index:1000;padding:.7rem 0;background:rgba(7,7,15,.55);backdrop-filter:blur(28px);-webkit-backdrop-filter:blur(28px);border-bottom:1px solid var(--bd)}}
.ni{{max-width:1160px;margin:0 auto;padding:0 2rem;display:flex;align-items:center;justify-content:space-between}}
.nl{{font-weight:800;font-size:1.05rem;color:#fff;text-decoration:none;letter-spacing:-.3px}}.nl span{{color:var(--a)}}
.nls{{display:flex;gap:.2rem}}.nls a{{color:var(--m);text-decoration:none;font-size:.8rem;font-weight:600;padding:.38rem .85rem;border-radius:6px;transition:all .18s}}.nls a:hover,.nls a.active{{color:#fff;background:rgba(255,255,255,.07)}}
.nav-cta{{padding:.42rem 1.15rem;background:var(--a);color:#fff;border-radius:8px;text-decoration:none;font-size:.8rem;font-weight:700;transition:all .18s}}.nav-cta:hover{{background:#7c7ff3;transform:translateY(-1px)}}
.wrap{{max-width:1160px;margin:0 auto;padding:0 2rem;position:relative;z-index:1}}
section{{padding:7rem 0}}
.sl{{font-size:.7rem;color:var(--a);letter-spacing:2.5px;text-transform:uppercase;font-weight:700;margin-bottom:.5rem}}
.st{{font-size:clamp(1.8rem,3.5vw,2.5rem);font-weight:800;color:#fff;letter-spacing:-1px;margin-bottom:2.8rem}}
/* Hero */
#hero{{min-height:100vh;display:flex;align-items:center;padding-top:4.5rem}}
.hg{{display:grid;grid-template-columns:1fr 360px;gap:5rem;align-items:center}}
.hey{{display:inline-flex;align-items:center;gap:.45rem;padding:.32rem .95rem;background:rgba(99,102,241,.1);border:1px solid rgba(99,102,241,.25);border-radius:50px;color:#a5b4fc;font-size:.72rem;font-weight:700;letter-spacing:1px;text-transform:uppercase;margin-bottom:1.4rem}}
.hey::before{{content:"";width:6px;height:6px;border-radius:50%;background:var(--a);box-shadow:0 0 10px var(--a);animation:pulse 2s infinite}}
@keyframes pulse{{0%,100%{{opacity:1;box-shadow:0 0 10px var(--a)}}50%{{opacity:.6;box-shadow:0 0 20px var(--a)}}}}
h1.hn{{font-size:clamp(3rem,6vw,5rem);font-weight:800;line-height:1.05;letter-spacing:-2.5px;color:#fff;margin-bottom:1.15rem}}
.grad{{background:linear-gradient(135deg,var(--a) 0%,var(--a2) 50%,var(--a4) 100%);-webkit-background-clip:text;-webkit-text-fill-color:transparent;background-clip:text;background-size:200%;animation:gs 4s ease infinite alternate}}
@keyframes gs{{0%{{background-position:0%}}100%{{background-position:100%}}}}
.hbio{{color:var(--ms);font-size:1.02rem;line-height:1.9;margin-bottom:2.4rem;max-width:490px}}
.hact{{display:flex;gap:.85rem;flex-wrap:wrap}}
.btn-p{{padding:.82rem 2rem;background:linear-gradient(135deg,var(--a),#7c7ff3);color:#fff;border-radius:10px;text-decoration:none;font-weight:700;font-size:.88rem;transition:all .22s;box-shadow:0 4px 24px rgba(99,102,241,.28);display:inline-flex;align-items:center;gap:.4rem}}.btn-p:hover{{transform:translateY(-2px);box-shadow:0 8px 32px rgba(99,102,241,.42)}}
.btn-o{{padding:.82rem 2rem;background:transparent;color:var(--t);border:1px solid var(--bd2);border-radius:10px;text-decoration:none;font-weight:600;font-size:.88rem;transition:all .22s;display:inline-flex;align-items:center;gap:.4rem}}.btn-o:hover{{background:var(--card);border-color:rgba(255,255,255,.2);transform:translateY(-2px)}}
/* Profile card */
.pfc{{background:var(--card);border:1px solid var(--bd);border-radius:var(--r2);padding:1.8rem;position:relative;overflow:hidden}}
.pfc::before{{content:"";position:absolute;top:0;left:0;right:0;height:3px;background:linear-gradient(90deg,var(--a),var(--a2),var(--a4))}}
.pf-av{{width:68px;height:68px;border-radius:50%;background:linear-gradient(135deg,var(--a),var(--a2));display:flex;align-items:center;justify-content:center;font-size:1.5rem;font-weight:800;color:#fff;margin-bottom:1rem;box-shadow:0 0 0 4px rgba(99,102,241,.15)}}
.pf-name{{font-weight:800;font-size:1.05rem;color:#fff;margin-bottom:.15rem}}
.pf-title{{font-size:.8rem;color:var(--a);font-weight:600;margin-bottom:1.4rem}}
.sg{{display:grid;grid-template-columns:1fr 1fr;gap:.5rem}}
.sb2{{background:var(--s1);border:1px solid var(--bd);border-radius:10px;padding:.8rem;text-align:center}}
.sn{{font-size:1.4rem;font-weight:800;color:#fff;line-height:1;margin-bottom:.12rem}}
.slb{{font-size:.62rem;color:var(--m);text-transform:uppercase;letter-spacing:.5px;font-weight:600}}
.gh-stats{{display:flex;gap:.4rem;margin-top:.5rem}}
.gh-s{{flex:1;background:var(--s1);border:1px solid var(--bd);border-radius:8px;padding:.55rem;text-align:center}}
.gh-n{{display:block;font-weight:700;font-size:.9rem;color:#fff}}.gh-l{{font-size:.6rem;color:var(--m)}}
/* Skills */
#skills{{background:var(--s1)}}
.sg-row{{display:flex;align-items:flex-start;gap:1.2rem;padding:.9rem 0;border-bottom:1px solid var(--bd)}}.sg-row:last-child{{border-bottom:none}}
.sg-cat{{font-size:.7rem;color:var(--m);text-transform:uppercase;letter-spacing:1px;font-weight:700;width:108px;flex-shrink:0;padding-top:.3rem}}
.sg-pills{{display:flex;flex-wrap:wrap;gap:.35rem}}
.sp{{padding:.28rem .8rem;background:var(--card);border:1px solid var(--bd);border-radius:6px;font-size:.78rem;color:#c0c0e0;font-weight:500;cursor:default;transition:all .18s}}.sp:hover{{border-color:rgba(99,102,241,.45);color:#fff;background:rgba(99,102,241,.08);transform:translateY(-1px)}}
/* Timeline */
.tl-item{{display:flex;gap:0;margin-bottom:1.8rem}}
.tl-lc{{display:flex;flex-direction:column;align-items:center;width:22px;flex-shrink:0;margin-top:4px}}
.tl-dot{{width:11px;height:11px;border-radius:50%;border:2px solid var(--m);background:var(--bg);flex-shrink:0;transition:all .3s}}
.tl-dot.current{{background:var(--a);border-color:var(--a);box-shadow:0 0 0 4px rgba(99,102,241,.18),0 0 14px rgba(99,102,241,.35)}}
.tl-line{{flex:1;width:2px;background:linear-gradient(to bottom,var(--bd),transparent);margin-top:.35rem}}
.tl-body{{flex:1;background:var(--card);border:1px solid var(--bd);border-radius:var(--r);padding:1.3rem 1.5rem;margin-left:.9rem;transition:all .22s}}.tl-body:hover{{border-color:rgba(99,102,241,.2);box-shadow:0 4px 20px rgba(0,0,0,.25)}}
.tl-hdr{{display:flex;justify-content:space-between;align-items:flex-start;gap:.8rem;margin-bottom:.4rem}}
.tl-role{{font-weight:700;font-size:.97rem;color:#fff}}.tl-company{{color:var(--a);font-size:.8rem;font-weight:600;margin-top:.18rem}}
.tl-dur{{font-size:.72rem;color:var(--m);white-space:nowrap;background:var(--s1);border:1px solid var(--bd);padding:.18rem .65rem;border-radius:50px;flex-shrink:0}}
.tl-desc{{font-size:.85rem;color:var(--ms);line-height:1.75;margin-top:.45rem}}
/* Projects */
#projects{{background:var(--s1)}}
.pj-grid{{display:grid;grid-template-columns:repeat(auto-fill,minmax(310px,1fr));gap:1.4rem}}
.pc{{position:relative;border-radius:var(--r2);overflow:hidden;border:1px solid var(--bd);transition:all .32s}}.pc:hover{{transform:translateY(-6px);border-color:rgba(99,102,241,.32);box-shadow:0 20px 50px rgba(0,0,0,.45)}}
.pc-glow{{position:absolute;inset:0;opacity:.55;transition:opacity .28s}}.pc:hover .pc-glow{{opacity:1}}
.pc-in{{position:relative;background:var(--card);margin:1px;border-radius:calc(var(--r2) - 1px);padding:1.7rem;display:flex;flex-direction:column;gap:.7rem;height:100%}}
.pc-top{{display:flex;align-items:center;justify-content:space-between}}
.pc-av{{width:42px;height:42px;border-radius:10px;display:flex;align-items:center;justify-content:center;font-size:1rem;font-weight:800;color:#fff}}
.pc-bgs{{display:flex;gap:.35rem;align-items:center}}
.pb-lang{{padding:.18rem .55rem;background:rgba(16,185,129,.08);border:1px solid rgba(16,185,129,.2);border-radius:4px;color:#6ee7b7;font-size:.66rem;font-weight:700}}
.pb-st{{color:#fbbf24;font-size:.72rem;font-weight:600}}
.pc-name{{font-weight:800;font-size:.97rem;color:#fff}}
.pc-desc{{font-size:.83rem;color:var(--ms);line-height:1.7;flex:1}}
.pc-topics{{display:flex;flex-wrap:wrap;gap:.28rem}}
.pt{{padding:.18rem .6rem;background:rgba(99,102,241,.08);border:1px solid rgba(99,102,241,.18);border-radius:4px;color:#a5b4fc;font-size:.66rem;font-weight:600}}
.pc-foot{{display:flex;align-items:center;justify-content:space-between;margin-top:.25rem}}
.pc-link{{color:var(--a);text-decoration:none;font-size:.8rem;font-weight:700;display:inline-flex;align-items:center;gap:.28rem;transition:all .18s}}.pc-link:hover{{color:#a5b4fc;gap:.55rem}}.pc-link span{{transition:transform .18s}}.pc-link:hover span{{transform:translateX(3px)}}
/* Education */
.edu-grid{{display:flex;flex-direction:column;gap:.9rem}}
.edu-card{{display:flex;align-items:center;gap:1.1rem;background:var(--card);border:1px solid var(--bd);border-radius:var(--r);padding:1.2rem 1.5rem;transition:all .22s}}.edu-card:hover{{border-color:rgba(99,102,241,.2)}}
.edu-icon{{font-size:1.4rem;flex-shrink:0}}
.edu-deg{{font-weight:700;color:#fff;font-size:.94rem}}.edu-inst{{color:var(--m);font-size:.8rem;margin-top:.22rem}}
.cert{{display:inline-block;margin:.28rem;padding:.32rem .95rem;background:rgba(245,158,11,.07);border:1px solid rgba(245,158,11,.2);border-radius:50px;color:#fcd34d;font-size:.76rem;font-weight:600}}
/* Contact */
#contact{{background:var(--s1);text-align:center}}
.ci{{max-width:540px;margin:0 auto}}
.c-av{{width:88px;height:88px;border-radius:50%;background:linear-gradient(135deg,var(--a),var(--a2));display:flex;align-items:center;justify-content:center;font-size:1.9rem;font-weight:800;color:#fff;margin:0 auto 1.4rem;box-shadow:0 0 0 8px rgba(99,102,241,.1),0 0 40px rgba(99,102,241,.18)}}
.ct{{font-size:2rem;font-weight:800;color:#fff;letter-spacing:-1px;margin-bottom:.7rem}}.cs{{color:var(--ms);line-height:1.85;margin-bottom:2.4rem}}
.sbs{{display:flex;justify-content:center;flex-wrap:wrap;gap:.7rem}}
.sb{{display:inline-flex;align-items:center;gap:.5rem;padding:.72rem 1.5rem;border-radius:10px;text-decoration:none;font-weight:700;font-size:.84rem;border:1px solid var(--bd);transition:all .22s}}.sb:hover{{transform:translateY(-2px)}}
.sb-li{{background:rgba(10,102,194,.1);color:#7ab8ff;border-color:rgba(10,102,194,.22)}}.sb-li:hover{{background:rgba(10,102,194,.18)}}
.sb-gh{{background:rgba(255,255,255,.05);color:#e0e0ff}}.sb-gh:hover{{background:rgba(255,255,255,.1)}}
.sb-em{{background:rgba(236,72,153,.08);color:#f9a8d4;border-color:rgba(236,72,153,.2)}}.sb-em:hover{{background:rgba(236,72,153,.14)}}
footer{{text-align:center;padding:2rem;color:var(--m);font-size:.76rem;border-top:1px solid var(--bd);position:relative;z-index:1}}
footer a{{color:var(--a);text-decoration:none}}
/* About contact */
.about-cl{{font-size:.7rem;color:var(--m);text-transform:uppercase;letter-spacing:1px;font-weight:700;margin-bottom:.4rem}}
.about-cv{{color:var(--ms);font-size:.88rem;margin-bottom:.28rem}}
/* Scroll reveal */
.rv{{opacity:0;transform:translateY(22px);transition:opacity .72s cubic-bezier(.16,1,.3,1),transform .72s cubic-bezier(.16,1,.3,1)}}.rv.in{{opacity:1;transform:none}}
.d1{{transition-delay:.06s}}.d2{{transition-delay:.14s}}.d3{{transition-delay:.22s}}
/* Hero fade in */
@keyframes fu{{from{{opacity:0;transform:translateY(28px)}}to{{opacity:1;transform:none}}}}
.f1{{animation:fu .88s cubic-bezier(.16,1,.3,1) .1s both}}.f2{{animation:fu .88s cubic-bezier(.16,1,.3,1) .22s both}}
.f3{{animation:fu .88s cubic-bezier(.16,1,.3,1) .34s both}}.f4{{animation:fu .88s cubic-bezier(.16,1,.3,1) .46s both}}
.f5{{animation:fu .88s cubic-bezier(.16,1,.3,1) .58s both}}
::-webkit-scrollbar{{width:3px}}::-webkit-scrollbar-track{{background:var(--bg)}}::-webkit-scrollbar-thumb{{background:rgba(99,102,241,.3);border-radius:3px}}
@media(max-width:860px){{
  .hg{{grid-template-columns:1fr;gap:2.5rem}}
  .pfc{{max-width:380px}}
  h1.hn{{font-size:clamp(2.4rem,8vw,3.5rem)}}
  .nls{{display:none}}
  section{{padding:5rem 0}}
  .sg-row{{flex-direction:column;gap:.4rem}}.sg-cat{{width:auto}}
  .pj-grid{{grid-template-columns:1fr}}
}}
</style>
</head>
<body>
<div class="orb o1"></div><div class="orb o2"></div><div class="orb o3"></div>

<nav>
  <div class="ni">
    <a href="#hero" class="nl">{initials[0]}<span>{initials[1] if len(initials)>1 else ""}</span></a>
    <div class="nls">
      <a href="#about">About</a><a href="#skills">Skills</a>
      <a href="#experience">Experience</a>{nav_projects}
      <a href="#education">Education</a><a href="#contact">Contact</a>
    </div>
    {hire_btn}
  </div>
</nav>

<section id="hero">
  <div class="wrap">
    <div class="hg">
      <div>
        <div class="hey f1">{title}</div>
        <h1 class="hn f2">Hi, I'm<br/><span class="grad">{name}</span></h1>
        <p class="hbio f3">{about[:220]}</p>
        <div class="hact f4">
          {"<a href='mailto:"+email+"' class='btn-p'>✉ Get In Touch</a>" if email else "<a href='#contact' class='btn-p'>Get In Touch</a>"}
          {"<a href='"+github+"' target='_blank' class='btn-o'>⚡ GitHub</a>" if github else "<a href='#projects' class='btn-o'>View Projects</a>"}
        </div>
      </div>
      <div class="pfc f5">
        <div class="pf-av">{initials}</div>
        <div class="pf-name">{name}</div>
        <div class="pf-title">{title}</div>
        <div class="sg">
          <div class="sb2"><div class="sn">{stat_num}</div><div class="slb">{stat_lbl}</div></div>
          <div class="sb2"><div class="sn">{len(skills)}</div><div class="slb">Technologies</div></div>
          <div class="sb2"><div class="sn">{len(projects) or len(exp)}</div><div class="slb">{"Projects" if projects else "Companies"}</div></div>
          <div class="sb2"><div class="sn">{len(certs) or len(edu)}</div><div class="slb">{"Certs" if certs else "Degrees"}</div></div>
        </div>
        {ghstats_h}
      </div>
    </div>
  </div>
</section>

<section id="about">
  <div class="wrap rv">
    <div class="sl">About Me</div><div class="st">The Story So Far</div>
    <div style="display:grid;grid-template-columns:1.2fr 1fr;gap:4rem;align-items:start">
      <p style="color:var(--ms);font-size:1.02rem;line-height:1.9">{about}</p>
      {about_contact}
    </div>
  </div>
</section>

<section id="skills">
  <div class="wrap">
    <div class="rv"><div class="sl">Expertise</div><div class="st">Technical Skills</div></div>
    <div class="rv d1">{skill_grid}</div>
  </div>
</section>

<section id="experience">
  <div class="wrap">
    <div class="rv"><div class="sl">Career</div><div class="st">Experience</div></div>
    {exp_html or "<div class='rv'><p style='color:var(--m)'>Experience details on request.</p></div>"}
  </div>
</section>

{"<section id='projects'><div class='wrap'><div class='rv'><div class='sl'>Work</div><div class='st'>Projects</div></div><div class='pj-grid'>"+proj_html+"</div></div></section>" if projects else ""}

<section id="education">
  <div class="wrap">
    <div class="rv"><div class="sl">Background</div><div class="st">Education</div></div>
    <div class="edu-grid">{edu_html or "<div class='edu-card rv'><div class='edu-icon'>🎓</div><div class='edu-info'><div class='edu-deg'>Education details on request.</div></div></div>"}</div>
    {"<div style='margin-top:2.2rem'><div class='rv'><p style='font-size:.7rem;color:var(--m);text-transform:uppercase;letter-spacing:1.5px;font-weight:700;margin-bottom:.9rem'>Certifications</p><div>"+certs_html+"</div></div></div>" if certs else ""}
  </div>
</section>

<section id="contact">
  <div class="wrap">
    <div class="ci rv">
      <div class="c-av">{initials}</div>
      <div class="ct">Let's Work Together</div>
      <p class="cs">I'm currently open to new opportunities. Whether you have a project, a question, or just want to say hi — my inbox is always open.</p>
      <div class="sbs">{soc}</div>
    </div>
  </div>
</section>

<footer><p>Built with <a href="https://github.com/Shweta-Mishra-ai/careerboost-ai">CareerBoost AI</a> · {name} © {year_now}</p></footer>

<script>
const obs=new IntersectionObserver(entries=>{{
  entries.forEach(e=>{{if(e.isIntersecting){{e.target.classList.add("in");obs.unobserve(e.target);}}}});
}},{{threshold:0.08,rootMargin:"0px 0px -35px 0px"}});
document.querySelectorAll(".rv").forEach(el=>obs.observe(el));
const secs=document.querySelectorAll("section[id]");
const navAs=document.querySelectorAll("nav a[href^='#']");
window.addEventListener("scroll",()=>{{
  let cur="";
  secs.forEach(s=>{{if(window.scrollY>=s.offsetTop-110)cur=s.id;}});
  navAs.forEach(a=>{{a.classList.toggle("active",a.getAttribute("href")==="#"+cur);}});
}},{{passive:true}});
</script>
</body>
</html>"""

    deploy_readme = """# How to Deploy Your Portfolio

## Option 1: Netlify (Easiest — 30 seconds)
1. Go to https://app.netlify.com/drop
2. Unzip this file → drag DROP the FOLDER (not zip) onto Netlify
3. Live instantly at a free URL

## Option 2: GitHub Pages
1. Create repo → upload index.html to root
2. Settings → Pages → Deploy from main → / (root)
3. Live at https://yourusername.github.io/reponame

## Option 3: Vercel
Drag the unzipped folder to vercel.com/new → Deploy
"""
    buf = io.BytesIO()
    with zipfile.ZipFile(buf, "w", zipfile.ZIP_DEFLATED) as zf:
        zf.writestr("index.html", html)
        zf.writestr("HOW_TO_DEPLOY.md", deploy_readme)
    buf.seek(0)
    return buf.read()
